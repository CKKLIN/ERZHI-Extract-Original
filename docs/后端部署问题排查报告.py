from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# ========== 封面标题 ==========
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ERZHI-Extract-Original\n后端部署问题排查报告')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'文档日期：2026年7月1日\n服务器：8.163.43.7 (CentOS 7)\n面板：宝塔Linux面板 v11.6.0')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ========== 目录 ==========
doc.add_heading('目录', level=1)
toc_items = [
    '1. 问题概述',
    '2. 排查过程',
    '   2.1 前端配置检查',
    '   2.2 服务器连通性测试',
    '   2.3 宝塔面板检查',
    '   2.4 防火墙规则检查',
    '   2.5 应用日志分析',
    '3. 根因分析',
    '4. 修复方案',
    '   4.1 代码修改',
    '   4.2 重新构建与部署',
    '5. 验证结果',
    '6. 当前状态',
    '7. 后续建议',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ========== 1. 问题概述 ==========
doc.add_heading('1. 问题概述', level=1)
doc.add_paragraph(
    '用户将后端 Spring Boot 应用部署到阿里云服务器（8.163.43.7）后，前端无法连接后端服务。'
    '前端配置已修改指向服务器地址，但所有 API 请求均返回连接失败。'
)

# ========== 2. 排查过程 ==========
doc.add_heading('2. 排查过程', level=1)

doc.add_heading('2.1 前端配置检查', level=2)
doc.add_paragraph(
    '检查 config/index.js 配置文件，确认 dev 和 prod 环境的 apiBase 均已修改为 http://8.163.43.7:3001。'
    '同时检查 manifest.json 中 H5 端的 devServer proxy 配置，发现代理目标仍指向 localhost:3001，'
    '已修改为服务器地址。前端配置确认无误。'
)

doc.add_heading('2.2 服务器连通性测试', level=2)
doc.add_paragraph(
    '从本地执行 curl 测试服务器 3001 端口：'
)
doc.add_paragraph(
    'curl http://8.163.43.7:3001/api/health\n'
    '结果：Connection refused（连接被拒绝）',
    style='List Bullet'
)
doc.add_paragraph('说明服务器端口未监听或服务未运行。')

doc.add_heading('2.3 宝塔面板检查', level=2)
doc.add_paragraph('登录宝塔面板（https://8.163.43.7:8888）后，进入「网站 → Java项目」页面，发现：')
doc.add_paragraph('Extract-Original 项目状态为「运行中」，但端口列显示 "--"，表示宝塔未能检测到应用监听的端口', style='List Bullet')
doc.add_paragraph('同一服务器上另一个 Java 项目 erzhi-server-0 运行正常，端口 8026 正常检测', style='List Bullet')
doc.add_paragraph('进程 PID 不断变化（8696 → 9054 → 9933），说明应用在不断重启', style='List Bullet')

doc.add_heading('2.4 防火墙规则检查', level=2)
doc.add_paragraph('检查宝塔「安全 → 系统防火墙」页面：')
doc.add_paragraph('端口 3001 规则已存在：tcp 3001 放行 入站 所有IP', style='List Bullet')
doc.add_paragraph('状态显示「未使用」，说明防火墙规则正确，但端口上无服务监听', style='List Bullet')

doc.add_heading('2.5 应用日志分析（关键发现）', level=2)
doc.add_paragraph('查看 Extract-Original 项目的日志管理，发现应用启动流程如下：')

doc.add_paragraph('阶段 1 - 正常启动：')
doc.add_paragraph(
    'Spring Boot 2.6.13 启动，Profile: prod\n'
    'Tomcat initialized with port(s): 3001 (http)\n'
    'HikariPool 数据库连接池启动成功\n'
    'JPA/Hibernate 初始化完成',
    style='List Bullet'
)

doc.add_paragraph('阶段 2 - Playwright 初始化失败：')
doc.add_paragraph(
    'Starting Playwright browser...\n'
    '/tmp/playwright-java-xxx/node: /lib64/libm.so.6: version GLIBC_2.27 not found\n'
    '/tmp/playwright-java-xxx/node: /lib64/libstdc++.so.6: version GLIBCXX_3.4.20 not found\n'
    '/tmp/playwright-java-xxx/node: /lib64/libc.so.6: version GLIBC_2.28 not found\n'
    '...',
    style='List Bullet'
)

doc.add_paragraph('阶段 3 - 应用崩溃：')
doc.add_paragraph(
    'Exception encountered during context initialization\n'
    '→ PlaywrightService 初始化失败\n'
    '→ MovieSearchService 依赖注入失败\n'
    '→ MovieController 创建失败\n'
    '→ 整个 Spring 容器启动失败\n'
    '→ Application run failed',
    style='List Bullet'
)

# ========== 3. 根因分析 ==========
doc.add_heading('3. 根因分析', level=1)

doc.add_paragraph(
    '根本原因是 Playwright（浏览器自动化库）的兼容性问题：'
)

doc.add_paragraph(
    'Playwright Java 1.40.0 内嵌了一个 Node.js 运行时，该 Node.js 需要 GLIBC 2.25 ~ 2.28。',
    style='List Bullet'
)
doc.add_paragraph(
    '服务器操作系统为 CentOS 7，其 GLIBC 版本为 2.17，无法满足 Playwright 的依赖要求。',
    style='List Bullet'
)
doc.add_paragraph(
    'PlaywrightService 在 @PostConstruct 初始化阶段抛出异常，导致 Spring Bean 创建失败。',
    style='List Bullet'
)
doc.add_paragraph(
    '由于 MovieController、MovieSearchService、VideoService 都直接或间接依赖 PlaywrightService，'
    'Bean 创建失败的连锁反应导致整个 Spring 容器无法启动，应用虽然进程存在但端口未监听。',
    style='List Bullet'
)

# 依赖关系图
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '依赖链：PlaywrightService(失败)\n'
    '  → MovieSearchService(失败)\n'
    '    → MovieController(失败)\n'
    '  → VideoService(失败)\n'
    '→ Spring 容器启动失败 → 端口 3001 未监听'
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_page_break()

# ========== 4. 修复方案 ==========
doc.add_heading('4. 修复方案', level=1)

doc.add_paragraph(
    '修改 PlaywrightService 使其在初始化失败时不阻塞整个应用的启动，'
    '而是降级处理：记录警告日志，将自身标记为不可用，让其他依赖 Playwright 的功能自动跳过或返回友好错误。'
)

doc.add_heading('4.1 代码修改', level=2)

doc.add_heading('文件：PlaywrightService.java', level=3)
doc.add_paragraph('路径：src/main/java/com/example/erzhiextractoriginalbackend/service/PlaywrightService.java')
doc.add_paragraph('修改内容：')

code_text = '''@Component
public class PlaywrightService implements DisposableBean {
    // 新增：可用性标志
    private boolean available = false;

    @PostConstruct
    public void init() {
        try {
            // 原有 Playwright 初始化逻辑
            log.info("Starting Playwright browser...");
            playwright = Playwright.create();
            // ... 浏览器启动配置 ...
            browser = playwright.chromium().launch(options);
            available = true;
            log.info("Browser started");
        } catch (Exception e) {
            // 新增：捕获异常，允许应用继续启动
            log.error("Failed to start Playwright browser (GLIBC may be too old): {}", e.getMessage());
            log.warn("Playwright-dependent features (Weixin, Movie search) will be unavailable");
            available = false;
        }
    }

    // 新增：检查 Playwright 是否可用
    public boolean isAvailable() {
        return available && browser != null;
    }
}'''

p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)
p.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()
doc.add_paragraph('MovieSearchService 已有容错处理：')
doc.add_paragraph('@Autowired(required = false) 注入 PlaywrightService', style='List Bullet')
doc.add_paragraph('search() 方法接受 Browser 参数，为 null 时自动跳过浏览器相关搜索（Bing、YouTube），仅返回 B站 API 结果', style='List Bullet')

doc.add_heading('4.2 重新构建与部署', level=2)
doc.add_paragraph('构建命令：mvn clean package -DskipTests')
doc.add_paragraph('构建产物：ERZHI-Extract-Original-Back-End-0.0.1-SNAPSHOT.jar（211MB）')
doc.add_paragraph('部署路径：/www/wwwroot/spring/ERZHI-Extract-Original-Back-End-0.0.1-SNAPSHOT.jar')
doc.add_paragraph('部署方式：通过宝塔面板文件管理器上传并覆盖旧 JAR，项目守护进程自动重启')

doc.add_page_break()

# ========== 5. 验证结果 ==========
doc.add_heading('5. 验证结果', level=1)

doc.add_paragraph('部署新 JAR 后，项目状态变化：')

table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '指标'
hdr_cells[1].text = '修复前'
hdr_cells[2].text = '修复后'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ('端口检测', '--（未检测到）', '3001 ✅'),
    ('API Health', 'Connection Refused', '{"success": true, "status": "ok"} ✅'),
    ('平台列表 API', 'Connection Refused', '返回 5 个平台 ✅'),
]
for i, (col1, col2, col3) in enumerate(data):
    row = table.rows[i + 1]
    row.cells[0].text = col1
    row.cells[1].text = col2
    row.cells[2].text = col3

doc.add_paragraph()

# ========== 6. 当前状态 ==========
doc.add_heading('6. 当前状态', level=1)

doc.add_paragraph('✅ 后端服务正常运行在 http://8.163.43.7:3001', style='List Bullet')
doc.add_paragraph('✅ 前端配置已指向服务器地址', style='List Bullet')
doc.add_paragraph('✅ 数据库连接正常（MySQL 3306）', style='List Bullet')
doc.add_paragraph('⚠️ Playwright 浏览器自动化功能不可用（受影响的模块见下方）', style='List Bullet')

doc.add_heading('Playwright 不可用影响的功能：', level=3)
doc.add_paragraph('微信视频号解析（可通过桥接代理降级，影响较小）', style='List Bullet')
doc.add_paragraph('影视搜索 - Bing 网页搜索', style='List Bullet')
doc.add_paragraph('影视搜索 - YouTube 搜索', style='List Bullet')
doc.add_paragraph('B站 搜索失败时的 Playwright 降级方案', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('不受影响的功能：抖音解析、小红书解析、B站解析、快手解析、YouTube解析、B站 API 搜索', style='List Bullet')

# ========== 7. 后续建议 ==========
doc.add_heading('7. 后续建议', level=1)

doc.add_paragraph('方案一：升级操作系统（推荐）')
doc.add_paragraph('将 CentOS 7 升级到 CentOS 8/9 或 Rocky Linux 8/9，原生支持 GLIBC 2.28+', style='List Bullet')
doc.add_paragraph('优点：彻底解决问题，Playwright 全部功能可用', style='List Bullet 2')
doc.add_paragraph('缺点：升级有风险，需要备份数据和配置', style='List Bullet 2')

doc.add_paragraph('方案二：使用 Docker 运行 Playwright')
doc.add_paragraph('将 Playwright 相关功能拆分到 Docker 容器中运行，容器内使用新版 Linux', style='List Bullet')
doc.add_paragraph('优点：不影响宿主机系统', style='List Bullet 2')
doc.add_paragraph('缺点：架构复杂度增加', style='List Bullet 2')

doc.add_paragraph('方案三：使用 Playwright 远程服务')
doc.add_paragraph('在另一台新版 Linux 服务器或云函数上运行 Playwright，通过 HTTP API 调用', style='List Bullet')
doc.add_paragraph('优点：完全解耦', style='List Bullet 2')
doc.add_paragraph('缺点：增加网络延迟和运维成本', style='List Bullet 2')

doc.add_paragraph('方案四：替换 Playwright 为 Selenium + 系统 Chrome')
doc.add_paragraph('使用 Selenium WebDriver + 系统安装的 Chrome/Chromium，不依赖 Playwright 内嵌的 Node.js', style='List Bullet')
doc.add_paragraph('优点：兼容 CentOS 7', style='List Bullet 2')
doc.add_paragraph('缺点：需要重写浏览器自动化代码', style='List Bullet 2')

# ========== 附录 ==========
doc.add_heading('附录：服务器环境信息', level=1)
doc.add_paragraph(f'操作系统：CentOS 7')
doc.add_paragraph(f'面板：宝塔Linux面板 v11.6.0')
doc.add_paragraph(f'JDK：jdk-17.0.8')
doc.add_paragraph(f'Web服务器：Nginx 1.28.1')
doc.add_paragraph(f'数据库：MySQL (8.163.43.7:3306)')
doc.add_paragraph(f'GLIBC 版本：2.17')
doc.add_paragraph(f'Playwright 版本：Java 1.40.0')
doc.add_paragraph(f'Spring Boot 版本：2.6.13')

# 保存文档
output_path = r'd:\money\ERZHI-Extract-Original\docs\ERZHI-Extract-Original后端部署问题排查报告.docx'
doc.save(output_path)
print(f'文档已保存到：{output_path}')
